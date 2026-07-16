"""
generar_doc.py — Autogenerador de la documentación del Integrador 2 en Word.

Lee los archivos Markdown de gobernanza del proyecto (catalogo-servicios.md,
matriz-resiliencia.md, matriz-auditoria.md, runbook.md) y produce un documento
Word profesional: Documentacion_Integrador2.docx

Requisito:
    pip install python-docx

Uso:
    python generar_doc.py
"""
import os
import re
import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAIZ = os.path.dirname(os.path.abspath(__file__))
SALIDA = os.path.join(RAIZ, "Documentacion_Integrador2.docx")


# ----------------------------------------------------------------------
# Helpers de Word
# ----------------------------------------------------------------------
def limpiar_inline(texto: str) -> str:
    """Quita marcado inline de Markdown (negritas, código, enlaces, emojis de header)."""
    texto = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", texto)  # [texto](url) -> texto
    texto = texto.replace("**", "").replace("`", "")
    return texto.strip()


def es_separador(celdas) -> bool:
    return all(set(c.strip()) <= set("-: ") and c.strip() for c in celdas)


def agregar_tabla(doc, filas_md):
    filas = []
    for linea in filas_md:
        celdas = [c.strip() for c in linea.strip().strip("|").split("|")]
        if es_separador(celdas):
            continue
        filas.append(celdas)
    if not filas:
        return
    ncols = max(len(f) for f in filas)
    tabla = doc.add_table(rows=0, cols=ncols)
    try:
        tabla.style = "Light Grid Accent 1"
    except KeyError:
        tabla.style = "Table Grid"
    for idx, fila in enumerate(filas):
        celdas_ui = tabla.add_row().cells
        for c in range(ncols):
            celdas_ui[c].text = limpiar_inline(fila[c]) if c < len(fila) else ""
        if idx == 0:  # cabecera en negrita
            for celda in celdas_ui:
                for par in celda.paragraphs:
                    for run in par.runs:
                        run.bold = True


def agregar_markdown(doc, md: str, offset: int = 1):
    """Vuelca contenido Markdown al documento (headings desplazados por `offset`)."""
    lineas = md.split("\n")
    i = 0
    while i < len(lineas):
        linea = lineas[i].rstrip()

        if linea.startswith("#"):
            m = re.match(r"(#+)\s+(.*)", linea)
            if m:
                nivel = min(len(m.group(1)) + offset, 4)
                doc.add_heading(limpiar_inline(m.group(2)), level=nivel)
        elif linea.startswith("```"):
            i += 1
            codigo = []
            while i < len(lineas) and not lineas[i].startswith("```"):
                codigo.append(lineas[i])
                i += 1
            par = doc.add_paragraph()
            run = par.add_run("\n".join(codigo))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif linea.strip().startswith("|") and "|" in linea.strip()[1:]:
            bloque = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                bloque.append(lineas[i])
                i += 1
            i -= 1
            agregar_tabla(doc, bloque)
        elif linea.startswith("> "):
            par = doc.add_paragraph(limpiar_inline(linea[2:]))
            try:
                par.style = "Quote"
            except KeyError:
                pass
        elif linea.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(limpiar_inline(linea.lstrip()[2:]), style="List Bullet")
        elif linea.strip():
            doc.add_paragraph(limpiar_inline(linea))
        i += 1


def leer_md(nombre: str) -> str:
    ruta = os.path.join(RAIZ, nombre)
    if not os.path.exists(ruta):
        return ""
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def agregar_toc(doc):
    """Inserta un índice automático de Word (se actualiza con F9 / clic derecho)."""
    par = doc.add_paragraph()
    run = par.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    texto = OxmlElement("w:t")
    texto.text = "Haz clic derecho aquí y elige 'Actualizar campo' para generar el índice."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, texto, end):
        run._r.append(el)


def portada(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SHServices V2")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Sistema de Help Desk & Punto de Venta (POS)\nArquitectura de Microservicios")
    rs.font.size = Pt(16)

    for _ in range(8):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Documentación Técnica — Curso Integrador II\n"
        f"Generado automáticamente el {datetime.date.today().strftime('%d/%m/%Y')}"
    ).font.size = Pt(12)
    doc.add_page_break()


# ----------------------------------------------------------------------
# Secciones escritas (contenido propio + embebido de los .md)
# ----------------------------------------------------------------------
def seccion_observabilidad(doc):
    doc.add_heading("Observabilidad", level=1)
    doc.add_paragraph(
        "La plataforma centraliza métricas y logs para operar los microservicios con visibilidad total:"
    )
    doc.add_paragraph("Prometheus: recolecta métricas (endpoint /metrics de cada servicio).", style="List Bullet")
    doc.add_paragraph("Grafana Loki: agrega los logs de TODOS los contenedores.", style="List Bullet")
    doc.add_paragraph("Promtail: descubre los contenedores vía el socket de Docker y envía sus logs a Loki.", style="List Bullet")
    doc.add_paragraph("Grafana (:3000): dashboards y consulta de logs (datasource Loki provisionado automáticamente).", style="List Bullet")
    doc.add_paragraph(
        "Cada petición viaja con un X-Correlation-ID generado por el Gateway, que se propaga por "
        "los servicios y los eventos de RabbitMQ, permitiendo reconstruir el rastro completo de una operación."
    )


def seccion_base_datos(doc):
    doc.add_heading("Base de Datos", level=1)
    doc.add_paragraph(
        "Persistencia en PostgreSQL (una base compartida, shservices_db). Cada microservicio con estado "
        "gestiona sus propias tablas mediante SQLAlchemy:"
    )
    doc.add_paragraph("tickets y garantias (ticket_service): ciclo de vida de atenciones y garantías de 90 días.", style="List Bullet")
    doc.add_paragraph("inventario (almacen_service): stock en 2 fases (disponible/reservado) y precio de venta.", style="List Bullet")
    doc.add_paragraph("diagnosticos (diagnostico_service): falla, mano de obra y repuestos (JSON).", style="List Bullet")
    doc.add_paragraph("facturas (facturacion_service): comprobantes con líneas de detalle.", style="List Bullet")
    doc.add_paragraph("usuarios (auth_service): empleados con rol y sede, con seed automático.", style="List Bullet")
    doc.add_paragraph("auditoria_eventos y notificaciones: trazabilidad y alertas internas.", style="List Bullet")
    doc.add_paragraph(
        "Las migraciones se realizan de forma NO destructiva con 'ALTER TABLE ... ADD COLUMN IF NOT EXISTS' "
        "al arranque de cada servicio, garantizando que los datos existentes nunca se pierden."
    )


def seccion_flujos(doc):
    doc.add_heading("Flujos Principales del Negocio", level=1)

    doc.add_heading("Soporte Técnico", level=2)
    doc.add_paragraph("CAJA registra el ticket (SOPORTE → EN_COLA) e imprime el Ticket de Recepción.", style="List Bullet")
    doc.add_paragraph("TECNICO diagnostica: define repuestos (se RESERVAN en almacén) y mano de obra → DIAGNOSTICADO.", style="List Bullet")
    doc.add_paragraph("Se notifica a CAJA que el equipo está listo para cobro y entrega.", style="List Bullet")
    doc.add_paragraph("CAJA cobra: se CONFIRMA (consume) el stock, se emite la boleta y se genera la garantía de 90 días → ENTREGADO.", style="List Bullet")
    doc.add_paragraph("Si el cliente rechaza el presupuesto, el stock reservado se LIBERA → RECHAZADO.", style="List Bullet")

    doc.add_heading("Venta Directa", level=2)
    doc.add_paragraph("CAJA selecciona productos del almacén; el stock se descuenta de golpe y se emite la boleta con detalle.", style="List Bullet")

    doc.add_heading("Garantías", level=2)
    doc.add_paragraph(
        "Toda reparación entregada genera una garantía de 90 días exactos. Recepción y Administración "
        "pueden consultarla por DNI/RUC o N° de serie, viendo su vigencia (VIGENTE/VENCIDA), días restantes "
        "y el monto cobrado por la reparación."
    )


# ----------------------------------------------------------------------
# Construcción del documento
# ----------------------------------------------------------------------
def main():
    doc = Document()

    portada(doc)

    doc.add_heading("Índice", level=1)
    agregar_toc(doc)
    doc.add_page_break()

    secciones_md = [
        ("Arquitectura de Microservicios", "catalogo-servicios.md"),
        ("Resiliencia y Circuit Breaker", "matriz-resiliencia.md"),
    ]
    for titulo, archivo in secciones_md:
        doc.add_heading(titulo, level=1)
        contenido = leer_md(archivo)
        if contenido:
            agregar_markdown(doc, contenido, offset=1)
        else:
            doc.add_paragraph(f"(No se encontró {archivo} en la raíz del proyecto.)")
        doc.add_page_break()

    seccion_observabilidad(doc)
    doc.add_page_break()

    seccion_base_datos(doc)
    doc.add_page_break()

    doc.add_heading("Auditoría y Trazabilidad", level=1)
    auditoria = leer_md("matriz-auditoria.md")
    if auditoria:
        agregar_markdown(doc, auditoria, offset=1)
    doc.add_page_break()

    seccion_flujos(doc)
    doc.add_page_break()

    doc.add_heading("Operación (Runbook)", level=1)
    runbook = leer_md("runbook.md")
    if runbook:
        agregar_markdown(doc, runbook, offset=1)

    doc.save(SALIDA)
    print(f"OK -> {SALIDA}")


if __name__ == "__main__":
    main()
